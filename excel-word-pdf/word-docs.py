'''
The Python-Docx third-party module can read and write .docx Word files.
Open a Word file with docx.Document()
Access one of the Paragraph objects from the "paragraphs" member variable, which is a list of Paragraph objects.
Paragraph objects have a "text" member variable containing the text as a string value.
Paragraphs are composed of "runs".  The "runs" member variable of a Paragraph object contains a list of Run objects.
Run objects also have a "text" member variable.
Run objects have a "bold", "italic", and "underline" member variables which can be set to True or False.
Paragraph and run objects have a "style" member variable that can be set to one of Word's built-in styles.
Word files can be created by calling add_paragraph() and add_run() to append text content.
'''

import docx

d = docx.Document('demo.docx')
print(d)
print(d.paragraphs)
print(d.paragraphs[0])
print(d.paragraphs[0].text)
print(d.paragraphs[1])
print(d.paragraphs[1].text)

p = d.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
p.runs[0].bold = None
print(p.runs[3].italic)
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'
d.save('demo2.docx')

print(p.style)
p.style = 'Title'
d.save('demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is another paragraph')
d.save('demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
print(p.runs)
p.runs[1].bold = True
d.save('demo5.docx')


def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


print(get_text('demo.docx'))
